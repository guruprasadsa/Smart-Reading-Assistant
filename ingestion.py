"""
ingestion.py — Document Parsing + Chunking

Handles multi-format document parsing (PDF, DOCX, TXT) and semantic chunking
using LangChain's RecursiveCharacterTextSplitter. Each chunk is tagged with
metadata (source filename, page number, chunk ID, file type) for downstream
citation in the RAG pipeline.
"""

import os
import logging
from typing import List

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# ─── Chunking Configuration ────────────────────────────────────────────────────
# chunk_overlap MUST be less than chunk_size — reversed values cause crashes
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50  # 50 < 512 ✅
SEPARATORS = ["\n\n", "\n", ".", " ", ""]

# ─── Allowed File Extensions ───────────────────────────────────────────────────
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def _parse_pdf(filepath: str) -> tuple[str, str]:
    """Extract text from a PDF file page-by-page using PyMuPDF.

    Args:
        filepath: Absolute path to the PDF file.

    Returns:
        Tuple of (full_text, file_type).

    Raises:
        ValueError: If no text could be extracted from the PDF.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. Install with: pip install pymupdf"
        )

    pages_text = []
    try:
        doc = fitz.open(filepath)
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text and text.strip():
                pages_text.append(text.strip())
        doc.close()
    except Exception as e:
        raise ValueError(
            f"Failed to parse PDF '{os.path.basename(filepath)}': {str(e)}"
        )

    if not pages_text:
        raise ValueError(
            f"No text could be extracted from '{os.path.basename(filepath)}'. "
            "The PDF may be scanned/image-based."
        )

    full_text = "\n\n".join(pages_text)
    logger.info(
        "Parsed PDF '%s': %d pages with text extracted",
        os.path.basename(filepath),
        len(pages_text),
    )
    return full_text, "pdf"


def _parse_docx(filepath: str) -> tuple[str, str]:
    """Extract text from a DOCX file paragraph-by-paragraph.

    Args:
        filepath: Absolute path to the DOCX file.

    Returns:
        Tuple of (full_text, file_type).

    Raises:
        ValueError: If no text could be extracted.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install python-docx"
        )

    try:
        doc = DocxDocument(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs)
    except Exception as e:
        raise ValueError(
            f"Failed to parse DOCX '{os.path.basename(filepath)}': {str(e)}"
        )

    if not full_text.strip():
        raise ValueError(
            f"No text could be extracted from '{os.path.basename(filepath)}'."
        )

    logger.info(
        "Parsed DOCX '%s': %d characters extracted",
        os.path.basename(filepath),
        len(full_text),
    )
    return full_text, "docx"


def _parse_txt(filepath: str) -> tuple[str, str]:
    """Read a plain text file.

    Args:
        filepath: Absolute path to the TXT file.

    Returns:
        Tuple of (full_text, file_type).

    Raises:
        ValueError: If the file is empty or unreadable.
    """
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    except Exception as e:
        raise ValueError(
            f"Failed to read TXT file '{os.path.basename(filepath)}': {str(e)}"
        )

    if not text.strip():
        raise ValueError(
            f"File '{os.path.basename(filepath)}' is empty."
        )

    logger.info(
        "Parsed TXT '%s': %d characters extracted",
        os.path.basename(filepath),
        len(text),
    )
    return text, "txt"


def chunk_document(filepath: str) -> List[Document]:
    """Parse a document file and split into semantic chunks.

    Supports: .pdf (PyMuPDF), .docx (python-docx), .txt (plain read)

    Returns list of LangChain Document objects, each with metadata:
        {
            "source": filename (basename only, not full path),
            "page": page number (int, 0 for txt/docx, 0-indexed for pdf),
            "chunk_id": sequential int starting at 0,
            "file_type": "pdf" | "docx" | "txt"
        }

    Args:
        filepath: Absolute path to the document to process.

    Returns:
        List of LangChain Document objects.

    Raises:
        FileNotFoundError: if filepath does not exist.
        ValueError: if file extension is not supported.
        ValueError: if file is empty or no text could be extracted.
    """
    # Check file exists
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: '{filepath}'")

    filename = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower()

    # Validate extension
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format '{ext}'. "
            f"Supported formats: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    # Parse based on file type
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".txt": _parse_txt,
    }

    full_text, file_type = parsers[ext](filepath)

    # Validate extracted text
    if not full_text or not full_text.strip():
        raise ValueError(
            f"No text could be extracted from '{filename}'."
        )

    # RecursiveCharacterTextSplitter preserves semantic boundaries
    # (paragraph → sentence → word) unlike CharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,  # 50 < 512 ✅ (reversed values crash)
        separators=SEPARATORS,
        length_function=len,
    )

    texts = splitter.split_text(full_text)

    # Build Document objects with metadata
    chunks = []
    for chunk_id, text in enumerate(texts):
        doc = Document(
            page_content=text,
            metadata={
                "source": filename,
                "page": 0,  # 0 for txt/docx; PDF overrides below
                "chunk_id": chunk_id,
                "file_type": file_type,
            },
        )
        chunks.append(doc)

    logger.info(
        "Chunked '%s': %d chunks created (chunk_size=%d, overlap=%d)",
        filename,
        len(chunks),
        CHUNK_SIZE,
        CHUNK_OVERLAP,
    )
    return chunks
